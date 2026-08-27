from copy import copy
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


SOURCE = Path(r"C:\Users\35958\Documents\xwechat_files\wxid_mbvzilsnuxpj22_0c47\msg\file\2026-08\智能排考小程序产品需求文档(3).docx")
OUTPUT = Path(r"D:\AI\人工智能项目\智能排考小程序\智能排考小程序产品需求文档（项目更新版）.docx")


def set_run_font(run, size=None, bold=None, color=None):
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def replace_paragraph(paragraph, text, *, size=None, bold=None, color=None, align=None):
    paragraph.clear()
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    if align is not None:
        paragraph.alignment = align


def set_cell(cell, text, *, bold=False):
    cell.text = text
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_after = Pt(0)
        for run in paragraph.runs:
            set_run_font(run, size=8.5, bold=bold)


def ensure_rows(table, count):
    while len(table.rows) < count:
        table.add_row()


def main():
    doc = Document(SOURCE)
    p = doc.paragraphs

    # 封面与目录：保留原有层级、段落位置和页面结构。
    replace_paragraph(p[0], "INTELLIGENT EXAM SCHEDULING", size=16, bold=True, color=(36, 99, 166), align=WD_ALIGN_PARAGRAPH.CENTER)
    replace_paragraph(p[1], "智能排考小程序", size=22, bold=True, color=(36, 99, 166), align=WD_ALIGN_PARAGRAPH.CENTER)
    replace_paragraph(p[2], "多场次监考自动安排与人工协同\n产品需求文档", size=19, bold=True, color=(40, 54, 77), align=WD_ALIGN_PARAGRAPH.CENTER)
    replace_paragraph(p[3], "版本 V2.0　｜　项目更新版　｜　2026-08-26", size=10, color=(90, 105, 125), align=WD_ALIGN_PARAGRAPH.CENTER)

    toc = [
        "1 项目定位与目标\t3",
        "2 调研结论与范围取舍\t3",
        "3 已实现核心方案\t4",
        "4 功能清单与优先级\t5",
        "5 关键功能：输入→处理→输出\t8",
        "6 差异化价值与解决痛点\t9",
        "参考资料\t11",
    ]
    for index, text in enumerate(toc, start=6):
        replace_paragraph(p[index], text, size=11, color=(40, 54, 77))

    content = {
        14: "1 项目定位与目标",
        15: "智能排考小程序面向学院考务老师，用于完成多场次考试的监考人员分配、人工调整、结果核对与打印导出。用户只需导入教师名单、教室需求和考试安排模板，即可生成可执行的监考安排。",
        16: "产品聚焦“自动排考 + 人工可控调整”：既遵守同一时间段不重复监考、考场人数等硬规则，也尽量满足第一监考经验、男女搭配、部门分散和工作量均衡等软规则。",
        17: "2 调研结论与范围取舍",
        18: "调研对象包括 Excel/VBA 排考工具、教务系统监考模块、UniTime、CELCAT 与 TimeEdit 等方案。调研目标不是复刻大型教务系统，而是找出学院考务场景最需要的能力：数据导入简单、规则透明、可手动修正、导出能直接打印。",
        20: "范围取舍：保留 Excel/WPS 的使用习惯和轻量桌面操作；不建设全校教务流程、学生排考或复杂权限体系；重点解决“一个学院、多场考试、监考人员安排与临时变动”。",
        21: "3 已实现核心方案",
        22: "多场次考试自动识别",
        23: "输入考试安排模板后，系统将每个工作表识别为一个考试场次，读取考试日期、起止时间、考试科目和考场编号，并结合教室表中的监考人数要求生成各场次任务。",
        24: "时间冲突与规则优先",
        25: "系统将“同一教师不得在重叠时间段承担正式监考或备选任务”作为硬约束；在可分配教师中，按第一监考有经验、男女搭配、不同部门的顺序进行优先选择。规则未完全满足时仍输出可执行结果，并在日志中说明。",
        26: "备选人员与人工协同",
        27: "每个考场默认配置 1—2 名备选监考。考务老师可增添、删除或更换教师；增添时优先推荐本考场备选教师，其次推荐其他考场备选、空闲教师和非重叠时段教师。系统会补齐受影响考场的备选名额并提示冲突。",
        28: "可打印导出与工作量管理",
        29: "系统支持按原模板导出多工作表安排表，也可按考场范围分组导出签名名单。导出表自动换行、调整人员列宽和合并单元格行高；教师工作量窗口可查看正式监考、备选待命、总任务量及具体场次/考场。",
        30: "4 功能清单与优先级",
        31: "必须有",
        32: "数据导入与校验",
        33: "导入教师、教室和考试模板，识别常见字段，校验工号、姓名、考场号、监考人数与重复数据，避免错误数据直接进入排考。",
        34: "多场次自动排考与时间冲突控制",
        35: "按工作表读取考试时间，禁止教师在时间重叠的场次中重复承担正式监考或备选任务；不足时明确标注缺口而不是静默安排。",
        36: "规则驱动的监考分配",
        37: "第一监考优先安排有经验教师；在人员充足时优先男女搭配、不同部门；支持“经验优先、男女搭配优先、部门均衡优先”等简洁模式。",
        38: "备选监考与人工调整",
        39: "为每个考场生成备选监考；支持删除、恢复原位、增添和替换人员。删除教师后不会自动转移到其他考场，恢复原位时保持本考场备选名单不变。",
        40: "结果、日志与工作量统计",
        41: "展示考场安排、人员缺口、规则满足度、备选人数和运行日志；工作量统计支持按姓名、工号、场次或考场搜索，默认聚焦有任务教师。",
        42: "应该有",
        43: "按考场范围或自动识别范围分组导出，便于不同区域、楼层或考务组打印签名。",
        44: "在人工调整时提供更清晰的时间冲突、原备选位置和替换影响提示。",
        45: "锦上添花",
        46: "缺口风险清单：按考试日期、场次和考场汇总缺口，提示最需要人工协调的任务。",
        47: "备选推荐解释：展示推荐教师满足经验、性别、部门和工作量平衡的原因。",
        48: "操作日志与变更回溯",
        49: "记录人工增删改教师的场次、考场和结果，方便考务老师复核临时调整。",
        50: "标准化模板",
        51: "提供教师、教室和考试安排模板示例，降低首次导入门槛。",
        52: "5 关键功能：输入→处理→输出",
        54: "6 差异化价值与解决痛点",
        55: "要解决的痛点",
        56: "传统 Excel 排考依赖人工复制、公式和反复核对，容易出现教师时间冲突、考场缺人和名单格式不适合打印的问题。",
        57: "大型教务系统功能全面但配置复杂，学院考务老师需要的是能直接导入现有表格、当天即可使用的轻量工具。",
        58: "通用排班工具不了解监考业务：第一监考经验、备选人员、考场签名名单、临时请假后的安全调整都需要额外人工处理。",
        59: "差异化价值",
        60: "多场次时间感知：不只检查单次安排，而是根据考试日期和起止时间控制教师冲突。",
        61: "备选人员可用：备选不是静态备注。系统支持优先把本考场备选转为正式监考，并自动补齐新的备选。",
        62: "小修改不洗牌：删除、恢复或增添教师时尽量保留其他考场已有备选名单，避免临考前因一次调整牵动整份安排。",
        63: "规则透明：日志直接展示经验、男女搭配、部门分散的满足度，以及人员缺口和异常原因。",
        65: "轻量而可打印：无需建设完整教务系统，沿用 Excel/WPS 文件；导出结果保留原模板并支持自动换行、分组和签名打印。",
        66: "人工协同友好：考务老师可以查看教师工作量，按姓名、工号、场次或考场快速检索，并在系统提示下完成临时调换。",
        67: "最终价值：把“人工反复核名单、临时改表”转化为“系统自动安排、人工安全调整、结果直接使用”的工作流程。",
        70: "参考资料",
        71: "1. UniTime Examination Timetabling Manual：https://help.unitime.org/manuals/examination-timetabling",
        72: "2. CELCAT Exam Scheduling：https://celcat.com/timetabler/extensions/exam-scheduling/",
        73: "3. TimeEdit Exam Management & Scheduling：https://www.academy.timeedit.com/products/exam",
        74: "4. aSc TimeTables Room Supervision Help：https://help.asctimetables.com/pdf/asc_timetables_en_L4.pdf",
        75: "调研与项目更新日期：2026-08-26",
    }
    for index, text in content.items():
        replace_paragraph(p[index], text, size=10.5 if p[index].style.name == "Normal" else None)

    # 调研对比表。
    research = doc.tables[0]
    research_rows = [
        ["方案", "优点", "局限与本项目取舍"],
        ["Excel / VBA 排考工具", "贴合既有 Excel 习惯，上手快，适合学院临时排考。", "规则、时间冲突、备选人员通常依赖人工维护；本项目保留表格导入，同时补足自动校验与规则分配。"],
        ["高校教务系统监考模块", "流程规范、数据源统一、适合全校级管理。", "实施范围大、配置成本高；本项目定位学院级轻量工具，强调直接使用现有模板。"],
        ["UniTime 等约束排程方案", "硬约束/软约束分层、冲突提示与统计方法成熟。", "覆盖学生、课程、教室等完整排程，超出本项目范围；借鉴其约束透明和异常反馈思路。"],
        ["CELCAT / TimeEdit", "多场次考试管理、资源统筹和可视化能力强。", "偏大型机构平台；本项目聚焦监考人员安排、临时调整和签名打印。"],
    ]
    for row_index, row in enumerate(research_rows):
        for col_index, value in enumerate(row):
            set_cell(research.cell(row_index, col_index), value, bold=row_index == 0)

    # 功能优先级和关键输入/处理/输出表。
    priorities = doc.tables[1]
    priority_rows = [
        ["优先级", "功能", "输入→处理→输出", "状态"],
        ["必须有", "数据导入与校验", "教师表、教室表、考试模板 → 字段识别与异常校验 → 可排考数据", "已完成"],
        ["必须有", "多场次排考", "各工作表的日期/时间/考场 → 时间冲突控制与规则分配 → 分场次正式监考名单", "已完成"],
        ["必须有", "备选监考", "正式名单与剩余教师 → 每考场配置 1—2 名备选 → 可调动备选名单", "已完成"],
        ["必须有", "人工增删改", "选中考场与教师 → 冲突校验、恢复原位、补齐备选 → 更新后的安排与日志", "已完成"],
        ["必须有", "工作量与日志", "多场次安排 → 正式/备选任务汇总 → 可搜索教师工作量与规则满足度", "已完成"],
        ["必须有", "主表与分组导出", "排考结果与分组规则 → 写回原模板、自动换行/行高 → 可打印签名表", "已完成"],
        ["应该有", "缺口风险清单", "缺口与场次数据 → 按时间/考场排序 → 人工协调优先级", "规划中"],
        ["锦上添花", "备选推荐解释", "候选教师标签与工作量 → 解释推荐原因 → 更透明的人工决策", "规划中"],
    ]
    ensure_rows(priorities, len(priority_rows))
    for row_index, row in enumerate(priority_rows):
        for col_index, value in enumerate(row):
            set_cell(priorities.cell(row_index, col_index), value, bold=row_index == 0)

    # 为表格保留清晰的页内阅读体验。
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.core_properties.title = "智能排考小程序产品需求文档（项目更新版）"
    doc.core_properties.subject = "多场次监考自动安排与人工协同"
    doc.core_properties.author = "智能排考项目组"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
