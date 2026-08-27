from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


SRC = Path(r"D:\AI\人工智能项目\智能排考小程序\智能排考小程序产品需求文档_融合优化版_目录美化版.docx")
OUT = Path(r"D:\AI\人工智能项目\智能排考小程序\智能排考小程序产品需求文档_融合优化版_章节精简版.docx")


def set_small_title_style(doc):
    try:
        style = doc.styles["小标题"]
    except KeyError:
        style = doc.styles.add_style("小标题", WD_STYLE_TYPE.PARAGRAPH)
    style.base_style = doc.styles["Normal"]
    style.font.name = "微软雅黑"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    style.font.size = Pt(11.5)
    style.font.bold = True
    style.font.color.rgb = RGBColor(31, 78, 121)
    style.paragraph_format.space_before = Pt(8)
    style.paragraph_format.space_after = Pt(3)
    style.paragraph_format.keep_with_next = True
    return style


def format_small_title(paragraph):
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.keep_with_next = True
    for run in paragraph.runs:
        run.font.name = "微软雅黑"
        run.font.size = Pt(11.5)
        run.font.bold = True
        run.font.color.rgb = RGBColor(31, 78, 121)


def add_toc_row(doc, anchor, text, page):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.tab_stops.add_tab_stop(Inches(6.35), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
    run = p.add_run(text)
    run.font.name = "宋体"
    run.font.size = Pt(11.5)
    run.font.bold = True
    run.font.color.rgb = RGBColor(35, 35, 35)
    p.add_run("\t")
    page_run = p.add_run(str(page))
    page_run.font.name = "宋体"
    page_run.font.size = Pt(11.5)
    page_run.font.color.rgb = RGBColor(35, 35, 35)
    anchor.addnext(p._p)
    return p._p


def remove_paragraph(paragraph):
    paragraph._p.getparent().remove(paragraph._p)


def main():
    OUT.write_bytes(SRC.read_bytes())
    doc = Document(str(OUT))
    small_title = set_small_title_style(doc)

    # Convert all subordinate heading levels into visual subheadings. They stay
    # easy to scan, but no longer create a deep chapter hierarchy or TOC clutter.
    for p in doc.paragraphs:
        if p.style.name in {"二级标题", "三级标题"}:
            p.style = small_title
            format_small_title(p)

    # Keep only major chapters in the directory.
    toc_title = next(p for p in doc.paragraphs if p.text.strip() == "目录")
    toc_title_index = next(i for i, p in enumerate(doc.paragraphs) if p._p is toc_title._p)
    next_page = next(
        p for p in doc.paragraphs[toc_title_index + 1 :]
        if p._p.xpath('.//*[local-name()="br" and @w:type="page"]')
    )
    next_page_index = next(i for i, p in enumerate(doc.paragraphs) if p._p is next_page._p)
    for p in list(doc.paragraphs[toc_title_index + 1 : next_page_index]):
        remove_paragraph(p)

    major_entries = [
        ("1．项目定位与范围", 3),
        ("2．同类方案调研与取舍", 3),
        ("3．核心规则与产品方案", 4),
        ("4．功能清单（按优先级）", 5),
        ("5．必须功能：输入→处理→输出", 8),
        ("6．差异化与待解决痛点", 8),
        ("7．范围边界与验收要点", 9),
        ("附：调研来源", 10),
    ]
    anchor = toc_title._p
    for text, page in major_entries:
        anchor = add_toc_row(doc, anchor, text, page)

    note = doc.add_paragraph("小标题用于阅读分层，不单独计入目录；页码按当前版本排版生成。")
    note.paragraph_format.space_before = Pt(8)
    note.paragraph_format.space_after = Pt(0)
    for run in note.runs:
        run.font.name = "宋体"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(110, 110, 110)
    anchor.addnext(note._p)

    # Avoid treating the directory title itself as a visible body chapter in
    # future automatic TOC updates.
    toc_title.style = doc.styles["Normal"]
    for run in toc_title.runs:
        run.font.name = "微软雅黑"
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = RGBColor(20, 20, 20)
    toc_title.alignment = 1

    doc.save(str(OUT))
    print(OUT)


if __name__ == "__main__":
    main()
