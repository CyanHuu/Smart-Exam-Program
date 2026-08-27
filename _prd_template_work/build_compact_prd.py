from datetime import date
from pathlib import Path
from shutil import copy2
import os
import tempfile
import zipfile
import xml.etree.ElementTree as ET

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


REFERENCE = Path(r"C:\Users\35958\Documents\xwechat_files\wxid_mbvzilsnuxpj22_0c47\msg\attach\5164728d9f2f58dab6623caae1850f9e\2026-08\Rec\3c9a26caeae7646b\F\5\AI淼淼\智启新程 - 最终版(2).docx")
COVER = Path(r"D:\AI\人工智能项目\智能排考小程序\智能排考封面背景.png")
OUT = Path(r"D:\AI\人工智能项目\智能排考小程序\智能排考小程序产品需求文档_精简版.docx")

BLUE = "0B55BC"
DARK = "063B87"
INK = "18212F"
MUTED = "667085"
SOFT = "EAF3FF"
PALE = "F5F8FC"
ORANGE = "F1A20A"
TABLE_WIDTH = 8316  # A4 minus 1.25 in side margins, in DXA


def set_run_font(run, size=10.5, color=INK, bold=None, italic=None):
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def clear_paragraph(p):
    for child in list(p._element):
        p._element.remove(child)


def set_paragraph(p, before=0, after=6, line=1.35, align=None, keep_next=False):
    f = p.paragraph_format
    f.space_before = Pt(before)
    f.space_after = Pt(after)
    f.line_spacing = line
    f.keep_with_next = keep_next
    if align is not None:
        p.alignment = align


def add_text(doc, text, style="Normal", after=6, color=INK, label=None):
    p = doc.add_paragraph(style=style)
    set_paragraph(p, after=after)
    if label and text.startswith(label):
        r = p.add_run(label)
        set_run_font(r, color=DARK, bold=True)
        r = p.add_run(text[len(label):])
        set_run_font(r, color=color)
    else:
        r = p.add_run(text)
        set_run_font(r, color=color)
    return p


def add_heading(doc, text, level):
    style = {1: "一级标题", 2: "二级标题", 3: "三级标题"}[level]
    p = doc.add_paragraph(style=style)
    set_paragraph(p, before={1: 16, 2: 10, 3: 6}[level], after={1: 8, 2: 5, 3: 3}[level], keep_next=True)
    r = p.add_run(text)
    set_run_font(r, size={1: 16, 2: 14, 3: 12}[level], color={1: DARK, 2: BLUE, 3: DARK}[level], bold=True)
    return p


def paragraph_left_rule(p, color=BLUE, width="18", space="8"):
    p_pr = p._p.get_or_add_pPr()
    border = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), width)
    left.set(qn("w:space"), space)
    left.set(qn("w:color"), color)
    border.append(left)
    p_pr.append(border)


def add_highlight(doc, title, text):
    p = doc.add_paragraph()
    set_paragraph(p, before=3, after=9, line=1.25)
    paragraph_left_rule(p)
    r = p.add_run(title + "  ")
    set_run_font(r, size=10.5, color=DARK, bold=True)
    r = p.add_run(text)
    set_run_font(r, size=10.5, color=INK)
    return p


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top=70, start=100, bottom=70, end=100):
    tc_pr = cell._tc.get_or_add_tcPr()
    mar = tc_pr.first_child_found_in("w:tcMar")
    if mar is None:
        mar = OxmlElement("w:tcMar")
        tc_pr.append(mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        element = mar.find(qn(f"w:{side}"))
        if element is None:
            element = OxmlElement(f"w:{side}")
            mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    ind = tbl_pr.first_child_found_in("w:tblInd")
    if ind is None:
        ind = OxmlElement("w:tblInd")
        tbl_pr.append(ind)
    ind.set(qn("w:w"), "120")
    ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    for grid_col, width in zip(table._tbl.tblGrid.gridCol_lst, widths):
        grid_col.set(qn("w:w"), str(width))
    for row_i, row in enumerate(table.rows):
        for col_i, cell in enumerate(row.cells):
            set_cell_width(cell, widths[col_i])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_i == 0:
                set_cell_shading(cell, SOFT)


def write_cell(cell, text, bold=False, color=INK, size=9.2):
    p = cell.paragraphs[0]
    clear_paragraph(p)
    set_paragraph(p, after=0, line=1.15)
    r = p.add_run(text)
    set_run_font(r, size=size, color=color, bold=bold)


def add_table(doc, headers, data, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    for i, text in enumerate(headers):
        write_cell(table.rows[0].cells[i], text, bold=True, color=DARK)
    for row in data:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            write_cell(cells[i], text)
    table_geometry(table, widths)
    p = doc.add_paragraph()
    set_paragraph(p, after=2)
    return table


def add_page_field(paragraph):
    run = paragraph.add_run()
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    run._r.append(field)
    set_run_font(run, size=9, color=MUTED)


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    rel = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    rpr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.append(underline)
    run.append(rpr)
    node = OxmlElement("w:t")
    node.text = text
    run.append(node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def reset_body(doc):
    body = doc._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def configure(doc):
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)
    section.header_distance = Inches(0.5)
    section.footer_distance = Inches(0.5)

    for style_name in ("Normal", "正文1"):
        style = doc.styles[style_name]
        style.font.name = "宋体"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        style.font.size = Pt(10.5 if style_name == "Normal" else 12)
        style.paragraph_format.line_spacing = 1.35
        style.paragraph_format.space_after = Pt(6)

    header = section.header
    for item in list(header._element):
        header._element.remove(item)
    p = header.add_paragraph()
    clear_paragraph(p)
    set_paragraph(p, after=0, align=WD_ALIGN_PARAGRAPH.RIGHT)
    r = p.add_run("智能排考小程序 · 产品需求文档")
    set_run_font(r, size=8.5, color=MUTED)

    footer = section.footer
    for item in list(footer._element):
        footer._element.remove(item)
    p = footer.add_paragraph()
    clear_paragraph(p)
    set_paragraph(p, after=0, align=WD_ALIGN_PARAGRAPH.CENTER)
    r = p.add_run("智能排考小程序   |   ")
    set_run_font(r, size=8.5, color=MUTED)
    add_page_field(p)


def prune_unused_media(path):
    """Remove the reference document's screenshots after its body was rewritten."""
    rel_ns = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
    main_ns = "http://schemas.openxmlformats.org/drawingml/2006/main"
    package_ns = "http://schemas.openxmlformats.org/package/2006/relationships"
    with zipfile.ZipFile(path, "r") as src:
        document_xml = ET.fromstring(src.read("word/document.xml"))
        used_ids = {
            node.attrib[f"{{{rel_ns}}}embed"]
            for node in document_xml.findall(f".//{{{main_ns}}}blip")
            if f"{{{rel_ns}}}embed" in node.attrib
        }
        rels_name = "word/_rels/document.xml.rels"
        rels = ET.fromstring(src.read(rels_name))
        unused_targets = set()
        for rel in list(rels):
            if rel.attrib.get("Type", "").endswith("/image") and rel.attrib.get("Id") not in used_ids:
                unused_targets.add("word/" + rel.attrib["Target"])
                rels.remove(rel)
        ET.register_namespace("", package_ns)
        rels_bytes = ET.tostring(rels, encoding="utf-8", xml_declaration=True)
        fd, temp_name = tempfile.mkstemp(suffix=".docx", dir=str(path.parent))
        os.close(fd)
        try:
            with zipfile.ZipFile(temp_name, "w", zipfile.ZIP_DEFLATED) as dst:
                for item in src.infolist():
                    if item.filename in unused_targets:
                        continue
                    payload = rels_bytes if item.filename == rels_name else src.read(item.filename)
                    dst.writestr(item, payload)
            src.close()
            os.replace(temp_name, path)
        finally:
            if os.path.exists(temp_name):
                os.unlink(temp_name)


def build():
    copy2(REFERENCE, OUT)
    doc = Document(OUT)
    reset_body(doc)
    configure(doc)

    # Cover page: derived from the reference's blue technology-report opening.
    p = doc.add_paragraph()
    set_paragraph(p, before=6, after=8)
    r = p.add_run("INTELLIGENT EXAM SCHEDULING")
    set_run_font(r, size=10, color=BLUE, bold=True)

    p = doc.add_paragraph()
    set_paragraph(p, after=4)
    r = p.add_run("智能排考小程序")
    set_run_font(r, size=27, color=DARK, bold=True)
    p = doc.add_paragraph()
    set_paragraph(p, after=16)
    r = p.add_run("监考人员自动安排与分组打印\n产品需求文档")
    set_run_font(r, size=15, color=INK, bold=True)

    pic_p = doc.add_paragraph()
    set_paragraph(pic_p, after=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    pic_p.add_run().add_picture(str(COVER), width=Inches(4.55))

    p = doc.add_paragraph()
    set_paragraph(p, after=3)
    r = p.add_run("面向学院考务老师的轻量化监考安排工具")
    set_run_font(r, size=11.5, color=DARK, bold=True)
    p = doc.add_paragraph()
    set_paragraph(p, after=1)
    r = p.add_run(f"版本 V1.1  |  需求确认日期 {date.today()}  |  内部评审稿")
    set_run_font(r, size=9.5, color=MUTED)
    doc.add_page_break()

    add_heading(doc, "1．项目概述", 1)
    add_text(doc, "智能排考小程序服务于学院考务老师的日常监考安排工作。考务老师导入教师名单和考场需求表后，系统根据“第一监考优先有经验、尽量男女搭配、尽量不同部门”的规则自动生成安排，并按考场号输出适合打印签名的分组名单。")
    add_highlight(doc, "核心边界", "本期仅处理同一批次的监考人员分配与打印交付；每位教师在一次排考中最多安排一个考场。教师不足时，系统必须显示缺口，不允许以重复安排掩盖风险。")
    add_heading(doc, "1.1 产品要解决的问题", 2)
    add_text(doc, "教师名单与考场表多以 Excel 形式保存，字段名称和格式不统一；人工排班需要反复核对经验、性别和部门，容易遗漏；生成总表后还需二次拆分打印。产品目标是把“导入、分配、核对、分组导出”收敛为一条可控工作流。")
    add_heading(doc, "1.2 产品目标", 2)
    add_text(doc, "在不改变学院既有 Excel 使用习惯的前提下，减少监考匹配和打印整理时间；让每一个不满足的规则、每一个人员缺口都明确可见；让最终名单能直接作为打印与签名材料使用。")

    add_heading(doc, "2．同类方案调研与启发", 1)
    add_text(doc, "调研以官方公开资料为依据。以下适配性判断围绕“学院级、Excel 驱动、监考人员分配”的已确认范围展开。", color=MUTED, after=5)
    add_table(doc, ["方案", "可借鉴点", "适配性判断"], [
        ("UniTime", "将硬约束与软偏好分层，并提供冲突和统计报告。", "方法论成熟，但覆盖考试时段、教室、学生等完整教务数据，实施范围远大于当前需求。"),
        ("CELCAT Exam Scheduling", "支持批量导入、复杂规则、结果预览和实时调整。", "适合大型院校；本项目更适合保留 Excel 导入和本地打印的轻量路径。"),
        ("TimeEdit Exam", "按监考人员的技能、资质、可用性和总工时安排。", "启发我们把“教师负荷可见”作为后续增强，而不是只关注能否填满考场。"),
        ("aSc Timetables", "自动生成巡视并按次数、分钟和教师约束进行优化。", "与本项目的“生成后人工核对”模式最接近，可借鉴公平分配与局部调整。"),
    ], [1600, 2920, 3796])
    add_highlight(doc, "调研结论", "不做覆盖全校教务的重型排考平台，而是吸收“规则分层、结果可解释、生成后可调整”的思想，服务考务老师最频繁的监考分配和分组打印任务。")

    add_heading(doc, "3．产品聚焦与规则设计", 1)
    add_heading(doc, "3.1 核心产品定位", 2)
    add_text(doc, "“随机但不失控，自动但可核对”的监考安排工具。系统先排除不可用情况，再在符合条件的教师中按优先级评分；同等条件下随机选择，降低固定排序带来的偏向。")
    add_heading(doc, "3.2 规则分层", 2)
    add_heading(doc, "3.2.1 必须遵守的硬约束", 3)
    add_text(doc, "教师工号唯一、考场编号有效、监考人数为正整数；每位教师在本次排考中只能被安排一次；教师不足时输出未完成考场和缺口数量。")
    add_heading(doc, "3.2.2 尽量满足的软规则", 3)
    add_text(doc, "第一监考优先有经验教师；当单场需要两名及以上监考时，尽量男女搭配；在此前提下尽量来自不同部门。任一软规则无法满足时，系统继续生成可执行方案，并在结果中标记原因。")
    add_heading(doc, "3.2.3 个性化增强方向", 3)
    add_text(doc, "后续可提供经验、性别、部门权重滑块及“经验优先、均衡优先、自定义”预设；同时显示每位教师本次安排情况，为学院逐步形成公平的安排依据。")

    add_heading(doc, "4．功能设计", 1)
    add_heading(doc, "4.1 必须有", 2)
    add_heading(doc, "4.1.1 教师名单导入与数据校验", 3)
    add_text(doc, "支持工号、姓名、性别、部门、是否有经验等字段的导入和常见中文列名识别。缺失工号、重复工号、性别或经验标签不合法时，系统定位错误行并阻止排程。")
    add_heading(doc, "4.1.2 考场需求导入与校验", 3)
    add_text(doc, "读取考场编号和所需监考人数；对空考场、重复考场和非正整数人数给出可操作的提示。")
    add_heading(doc, "4.1.3 自动监考分配", 3)
    add_text(doc, "系统先尝试为每个考场分配有经验的第一监考，再按性别搭配和部门差异补足其余人员；教师不重复使用；同分候选随机抽取。")
    add_heading(doc, "4.1.4 结果核对与异常提示", 3)
    add_text(doc, "结果页显示考场、第一监考、其余监考人员和规则状态。教师总数不足、经验教师不足、无法男女搭配或不同部门时，均需以醒目提示呈现。")
    add_heading(doc, "4.1.5 完整表与分组表导出", 3)
    add_text(doc, "在既有考场安排模板中写入监考人员，并支持按考场号区间或列表分组，例如“101-112；201,203,205”。导出文件保留标题行并预留签名、备注等打印区域。")
    add_heading(doc, "4.2 应该有", 2)
    add_text(doc, "规则权重滑块、教师安排次数统计、锁定已确认考场后局部重排、候补教师建议，以及常用 Excel 字段映射记忆。")
    add_heading(doc, "4.3 锦上添花", 2)
    add_text(doc, "排程版本对比、可配置院系抬头和签名栏、导入模板一键复用。")

    add_heading(doc, "5．必须功能的输入、处理与输出", 1)
    add_table(doc, ["功能", "输入", "系统处理", "输出"], [
        ("教师名单校验", "教师 Excel：工号、姓名、性别、部门、经验。", "识别字段、校验工号唯一性与标签值。", "标准化教师清单或错误清单。"),
        ("考场需求校验", "考场 Excel：考场编号、监考人数。", "校验编号非空、人数为正整数。", "可排程的考场需求清单。"),
        ("自动监考分配", "已校验的教师、考场和规则。", "硬约束过滤；经验、性别、部门依次评分；同分随机。", "考场—教师安排及第一监考标识。"),
        ("异常核对", "自动分配结果。", "统计人员缺口与软规则未满足情况。", "待人工处理的考场与原因。"),
        ("完整表导出", "确认结果和目标模板。", "按考场匹配对应行并写入姓名。", "完整监考安排 Excel。"),
        ("分组表导出", "完整表及考场号分组规则。", "解析范围/列表并复制标题、记录和签名区。", "可打印的分组工作表或文件。"),
    ], [1350, 1900, 2870, 2196])

    add_heading(doc, "6．差异化与产品价值", 1)
    add_heading(doc, "6.1 可解释的随机", 2)
    add_text(doc, "系统不只是给出名单，还说明第一监考的经验状态、性别搭配与部门差异是否满足。考务老师可以快速理解结果，而不是面对一个无法说明的随机表。")
    add_heading(doc, "6.2 不掩盖真实缺口", 2)
    add_text(doc, "教师不足时拒绝重复安排，把缺少的人数、未填满的考场直接暴露出来。它解决的是考务现场最危险的问题：一张看起来完整、实际上无法执行的表。")
    add_heading(doc, "6.3 为 Excel 和打印而设计", 2)
    add_text(doc, "无需迁移到重型教务系统。导入格式、结果表和分组打印均以考务老师的现有工作习惯为出发点，降低培训和落地成本。")

    add_heading(doc, "7．范围边界与验收建议", 1)
    add_text(doc, "本版本不负责考试日期和时间编排、学生座位、课程冲突、跨校区调度、教师线上确认或请假替班。若未来需要同一教师跨时段监考，必须增加日期与时段字段并重新定义冲突规则。")
    add_highlight(doc, "验收重点", "导入错误可定位；任何教师在结果中不重复出现；有经验教师充足时各场第一监考均有经验；教师不足时出现缺口而不是重复分配；按考场号分组导出的表可直接打印签名。")

    add_heading(doc, "附：调研来源", 1)
    sources = [
        ("UniTime — Examination Timetabling Manual", "https://help.unitime.org/manuals/examination-timetabling"),
        ("CELCAT — Exam Scheduling", "https://celcat.com/timetabler/extensions/exam-scheduling/"),
        ("TimeEdit — Exam Management & Scheduling", "https://www.academy.timeedit.com/products/exam"),
        ("aSc TimeTables — Room Supervision Help", "https://help.asctimetables.com/pdf/asc_timetables_en_L4.pdf"),
    ]
    for i, (name, url) in enumerate(sources, 1):
        p = doc.add_paragraph()
        set_paragraph(p, after=3, line=1.1)
        r = p.add_run(f"{i}. {name}：")
        set_run_font(r, size=9.2, color=INK)
        add_hyperlink(p, url, url)
    add_text(doc, "调研访问日期：2026-08-25。产品能力描述以官方公开页面为准；适配性评价基于本项目已确认的范围。", after=0, color=MUTED)

    doc.save(OUT)
    prune_unused_media(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
