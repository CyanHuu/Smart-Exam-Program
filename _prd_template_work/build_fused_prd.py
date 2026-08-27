from datetime import date
from pathlib import Path
from shutil import copy2
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

sys.path.insert(0, r"D:\AI\人工智能项目\智能排考小程序\_prd_template_work")
from build_compact_prd import (
    COVER,
    BLUE,
    DARK,
    INK,
    MUTED,
    TABLE_WIDTH,
    add_heading,
    add_highlight,
    add_hyperlink,
    add_table,
    add_text,
    configure,
    prune_unused_media,
    reset_body,
    set_paragraph,
    set_run_font,
)


SOURCE = Path(r"D:\AI\人工智能项目\智能排考小程序\智能排考小程序产品需求文档_精简版.docx")
OUT = Path(r"D:\AI\人工智能项目\智能排考小程序\智能排考小程序产品需求文档_融合优化版.docx")


def add_cover(doc):
    p = doc.add_paragraph()
    set_paragraph(p, before=6, after=8)
    r = p.add_run("INTELLIGENT EXAM SCHEDULING")
    set_run_font(r, size=10, color=BLUE, bold=True)

    p = doc.add_paragraph()
    set_paragraph(p, after=4)
    r = p.add_run("智能排考小程序")
    set_run_font(r, size=27, color=DARK, bold=True)
    p = doc.add_paragraph()
    set_paragraph(p, after=16)
    r = p.add_run("监考人员自动安排与分组打印\n产品需求文档（融合优化版）")
    set_run_font(r, size=15, color=INK, bold=True)
    picture = doc.add_paragraph()
    set_paragraph(picture, after=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    picture.add_run().add_picture(str(COVER), width=Inches(4.55))
    p = doc.add_paragraph()
    set_paragraph(p, after=3)
    r = p.add_run("面向学院考务老师的 Excel 驱动监考安排工具")
    set_run_font(r, size=11.5, color=DARK, bold=True)
    p = doc.add_paragraph()
    set_paragraph(p, after=1)
    r = p.add_run(f"版本 V1.2  |  融合优化日期 {date.today()}  |  内部评审稿")
    set_run_font(r, size=9.5, color=MUTED)
    doc.add_page_break()


def build():
    copy2(SOURCE, OUT)
    doc = Document(OUT)
    reset_body(doc)
    configure(doc)
    add_cover(doc)

    add_heading(doc, "1．项目定位与范围", 1)
    add_text(doc, "智能排考小程序面向学院考务老师，解决已有考场安排后的监考人员分配和分组打印问题。考务老师无需迁移原有 Excel，只需导入教师名单和考场需求，即可得到一张可核对、可导出、可打印的监考安排表。")
    add_highlight(doc, "本期聚焦", "监考人员名单导入、考场需求导入、按优先级自动分配、异常提示、完整表导出与按考场号分组导出。")
    add_text(doc, "本期不处理考试时间编排、学生座位、课程冲突、跨校区调度、教师线上确认或请假替班。这样可以把需求收敛到考务老师最频繁、最耗时的实际动作上。")

    add_heading(doc, "2．同类方案调研与取舍", 1)
    add_text(doc, "本次调研同时采用 WPS 需求文档中的“常见方案分类”与前期官方资料调研。比较目标不是寻找功能最多的系统，而是确定本项目应保留什么、避免什么。", color=MUTED, after=5)
    add_table(doc, ["方案", "值得保留", "应避免"], [
        ("Excel 函数 / VBA 排考工具", "沿用 Excel 导入和导出习惯，上手快，适合学院现有资料。", "复杂优先级容易变成难维护的公式；缺少规则解释和异常可视化。"),
        ("高校教务系统监考模块", "规则完整、表格格式规范、可统一管理。", "实施范围大、配置重，学院级的灵活分组打印往往不够顺手。"),
        ("通用人员排班小程序", "人员任务分配、基础随机和统计交互较轻量。", "不了解“第一监考经验优先”、考场人数和监考签名表等考务语义。"),
        ("UniTime 等成熟排程方案", "硬约束/软规则分层、冲突提示、结果统计的方法论。", "覆盖考试时段、教室、学生等全流程，超出本项目的必要范围。"),
    ], [1780, 3180, 3356])
    add_highlight(doc, "取舍结论", "保留 Excel 的低门槛、成熟排程的规则分层和通用排班的轻交互；去掉重型教务平台的全流程建设，也去掉只靠随机抽人的不可解释性。")

    add_heading(doc, "3．核心规则与产品方案", 1)
    add_heading(doc, "3.1 可执行性优先", 2)
    add_text(doc, "每位教师在一次排考中最多安排一个考场。教师总数不足时，系统不得重复安排教师；应明确显示缺口数量、未完成考场和需人工协调的原因。")
    add_heading(doc, "3.2 软规则按优先级执行", 2)
    add_text(doc, "在满足硬约束的候选教师中，系统按“第一监考有经验 > 男女搭配 > 不同部门”的顺序进行选择。经验、性别和部门均为尽量满足的偏好：某一条件无法满足时，系统应继续生成可执行结果，并解释妥协原因。")
    add_heading(doc, "3.3 随机必须可核对", 2)
    add_text(doc, "同等优先级候选人随机抽取，避免固定排序导致总是同一批教师被优先安排。每次生成结果记录生成时间；结果页同时显示第一监考、经验状态、性别搭配和部门差异，方便考务老师人工复核。")
    add_heading(doc, "3.4 分组以考场号为中心", 2)
    add_text(doc, "分组规则直接使用考场号范围或考场号列表，例如“101-112；201,203,205”，而不是要求用户额外维护序号。系统导出的分组表保留标题、监考人员、签名与备注区域，直接服务打印场景。")

    add_heading(doc, "4．功能清单（按优先级）", 1)
    add_heading(doc, "4.1 必须有", 2)
    add_heading(doc, "4.1.1 Excel 导入与数据校验", 3)
    add_text(doc, "导入教师工号、姓名、性别、部门、经验，以及考场编号、所需监考人数。系统识别常见列名，定位空值、重复工号、重复考场和非正整数人数。")
    add_heading(doc, "4.1.2 排考规则设置", 3)
    add_text(doc, "默认固化“经验 > 性别 > 部门”的规则顺序；明确区分硬约束与可放宽偏好，防止规则一多就无法排出结果。")
    add_heading(doc, "4.1.3 自动监考分配", 3)
    add_text(doc, "优先为各考场选择有经验的第一监考，再补足其他监考人员；尽量男女搭配、不同部门；同等候选随机选择；任何教师均不重复分配。")
    add_heading(doc, "4.1.4 结果核对与异常提示", 3)
    add_text(doc, "清晰展示每个考场的人员、第一监考和规则状态；对人员不足、经验教师不足、性别或部门偏好未满足等情况给出原因，而不是静默输出结果。")
    add_heading(doc, "4.1.5 总表与分组表导出", 3)
    add_text(doc, "按给定模板输出完整监考表，并按用户输入的考场号规则输出多个分组 Sheet 或文件，保留标题行和打印签名区域。")
    add_heading(doc, "4.2 应该有", 2)
    add_text(doc, "经验、性别、部门的权重滑块；教师本次是否被安排、所在考场和第一监考次数统计；重新生成方案；常用 Excel 列映射和导出模板记忆。")
    add_heading(doc, "4.3 锦上添花", 2)
    add_text(doc, "排程版本对比、院系抬头与签名栏样式配置、打印前预览。候补池、局部重排和线上通知不进入当前版本，待真实使用后再评估。")

    add_heading(doc, "5．必须功能：输入→处理→输出", 1)
    add_table(doc, ["功能", "输入", "处理", "输出"], [
        ("教师名单校验", "教师 Excel：工号、姓名、性别、部门、经验。", "字段匹配；校验工号唯一、标签有效。", "标准教师清单或带行号的错误提示。"),
        ("考场需求校验", "考场 Excel：考场号、监考人数。", "校验编号非空、人数为正整数。", "可排程的考场需求清单。"),
        ("自动监考分配", "已校验数据及默认优先级。", "先过滤硬约束；按经验、性别、部门评分；同分随机。", "考场—教师安排与第一监考标识。"),
        ("异常核对", "自动分配结果。", "统计人员缺口和软规则妥协。", "待人工协调的考场、缺口和原因。"),
        ("完整表导出", "确认结果和目标模板。", "按考场号写入监考人员，保留原表标题。", "完整监考安排 Excel。"),
        ("分组表导出", "完整表及考场号范围/列表。", "解析分组规则，复制对应记录与标题。", "可直接打印签名的分组表。"),
    ], [1350, 1900, 2870, 2196])

    add_heading(doc, "6．差异化与待解决痛点", 1)
    add_heading(doc, "6.1 面向考务，而非泛排班", 2)
    add_text(doc, "系统理解第一监考、考场所需人数、经验标签与签名打印这些考务概念；它不是把教师当成通用任务资源的排班工具。")
    add_heading(doc, "6.2 自动但不黑箱", 2)
    add_text(doc, "结果不仅给出谁被安排，还能指出哪些规则已满足、哪些因为候选不足而被放宽。这样能减少考务老师反复解释“为什么是这位教师”的成本。")
    add_heading(doc, "6.3 不用重复安排制造假象", 2)
    add_text(doc, "人员不足时保留真实问题，不把同一教师安排到两个不同考场。该约束保证导出的名单是可执行名单，而不是形式上的填表结果。")
    add_heading(doc, "6.4 适配原有 Excel 与打印流程", 2)
    add_text(doc, "从导入到按考场号分组导出均围绕现有资料和打印习惯设计，降低学院采用成本。")

    add_heading(doc, "7．范围边界与验收要点", 1)
    add_text(doc, "验收时必须验证：导入错误可定位；任何教师在结果中只出现一次；经验教师充足时每个考场的第一监考均有经验；人数不足时生成缺口而非重复安排；软规则未满足时可见原因；按考场号分组导出的表可直接打印。")
    add_highlight(doc, "已去除的内容", "不将候补教师、局部重排、考试时间/座位、跨校区调度和线上通知写入本期核心需求，避免初版功能过多、规则边界不清。")

    add_heading(doc, "附：调研来源", 1)
    sources = [
        ("用户提供：自动排考程序_需求文档（WPS 云文档）", "https://kdocs.cn/l/caVIcZ7dv3d9"),
        ("UniTime — Examination Timetabling Manual", "https://help.unitime.org/manuals/examination-timetabling"),
        ("CELCAT — Exam Scheduling", "https://celcat.com/timetabler/extensions/exam-scheduling/"),
        ("TimeEdit — Exam Management & Scheduling", "https://www.academy.timeedit.com/products/exam"),
        ("aSc TimeTables — Room Supervision Help", "https://help.asctimetables.com/pdf/asc_timetables_en_L4.pdf"),
    ]
    for i, (name, url) in enumerate(sources, 1):
        p = doc.add_paragraph()
        set_paragraph(p, after=3, line=1.1)
        r = p.add_run(f"{i}. {name}：")
        set_run_font(r, size=9.2, color=INK)
        add_hyperlink(p, url, url)
    add_text(doc, "调研访问日期：2026-08-25。官方产品能力以公开页面为准；本项目的取舍基于已确认的学院级监考安排范围。", after=0, color=MUTED)

    doc.save(OUT)
    prune_unused_media(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
