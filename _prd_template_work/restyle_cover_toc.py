from pathlib import Path

from docx import Document
from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


SRC = Path(r"D:\AI\人工智能项目\智能排考小程序\智能排考小程序产品需求文档_融合优化版_封面目录优化.docx")
OUT = Path(r"D:\AI\人工智能项目\智能排考小程序\智能排考小程序产品需求文档_融合优化版_目录美化版.docx")


def shade(paragraph, fill="06265F"):
    ppr = paragraph._p.get_or_add_pPr()
    old = ppr.find(qn("w:shd"))
    if old is not None:
        ppr.remove(old)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    ppr.append(shd)


def set_cover_run(run, size, bold=True):
    run.font.name = "微软雅黑"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(255, 255, 255)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "微软雅黑")


def remove_element(paragraph):
    paragraph._p.getparent().remove(paragraph._p)


def add_toc_row(doc, anchor, text, page, level):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.tab_stops.add_tab_stop(Inches(6.35), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
    p.paragraph_format.left_indent = Inches(0.08 * level)
    p.paragraph_format.first_line_indent = Inches(-0.08 * level)
    run = p.add_run(text)
    run.font.name = "宋体"
    run.font.size = Pt(11.5 if level == 0 else 10.5)
    run.font.bold = level == 0
    run.font.color.rgb = RGBColor(35, 35, 35)
    p.add_run("\t")
    page_run = p.add_run(str(page))
    page_run.font.name = "宋体"
    page_run.font.size = Pt(11.5 if level == 0 else 10.5)
    page_run.font.color.rgb = RGBColor(35, 35, 35)
    anchor.addnext(p._p)
    return p._p


def main():
    OUT.write_bytes(SRC.read_bytes())
    doc = Document(str(OUT))

    # Improve cover contrast: white typography on dark navy bands.
    cover_break = next(p for p in doc.paragraphs if p._p.xpath('.//*[local-name()="br" and @w:type="page"]'))
    cover_paras = []
    for p in doc.paragraphs:
        if p._p is cover_break._p:
            break
        cover_paras.append(p)
    sizes = [12, 28, 17, 11, 9]
    for idx, p in enumerate(cover_paras):
        if not p.text.strip():
            continue
        shade(p)
        p.paragraph_format.space_before = Pt(2 if idx else 0)
        p.paragraph_format.space_after = Pt(5 if idx < 3 else 2)
        for run in p.runs:
            set_cover_run(run, sizes[min(idx, len(sizes) - 1)], bold=idx != 4)

    # Replace the field fallback with a visual TOC matching the supplied sample.
    toc_title = next(p for p in doc.paragraphs if p.text.strip() == "目录")
    toc_field = next(p for p in doc.paragraphs if p._p.xpath('.//*[local-name()="fldChar"]'))
    note = next((p for p in doc.paragraphs if p.text.startswith("提示：在 Word")), None)
    page_break_after_toc = next(p for p in doc.paragraphs if p._p.xpath('.//*[local-name()="br" and @w:type="page"]') and p._p is not cover_break._p)
    remove_element(toc_field)
    if note is not None:
        remove_element(note)

    entries = [
        ("1．项目定位与范围", 3, 0),
        ("2．同类方案调研与取舍", 3, 0),
        ("3．核心规则与产品方案", 4, 0),
        ("3.1 可执行性优先", 4, 1),
        ("3.2 软规则按优先级执行", 4, 1),
        ("3.3 随机必须可核对", 5, 1),
        ("3.4 分组以考场号为中心", 5, 1),
        ("4．功能清单（按优先级）", 5, 0),
        ("4.1 必须有", 5, 1),
        ("4.1.1 Excel 导入与数据校验", 5, 2),
        ("4.1.2 排考规则设置", 6, 2),
        ("4.1.3 自动监考分配", 6, 2),
        ("4.1.4 结果核对与异常提示", 6, 2),
        ("4.1.5 总表与分组表导出", 6, 2),
        ("4.2 应该有", 7, 1),
        ("4.3 锦上添花", 7, 1),
        ("5．必须功能：输入→处理→输出", 8, 0),
        ("6．差异化与待解决痛点", 8, 0),
        ("6.1 面向考务，而非泛排班", 8, 1),
        ("6.2 自动但不黑箱", 8, 1),
        ("6.3 不用重复安排制造假象", 9, 1),
        ("6.4 适配原有 Excel 与打印流程", 9, 1),
        ("7．范围边界与验收要点", 9, 0),
        ("附：调研来源", 10, 0),
    ]
    anchor = toc_title._p
    for text, page, level in entries:
        anchor = add_toc_row(doc, anchor, text, page, level)

    # Keep the document self-explanatory without cluttering the directory.
    note = doc.add_paragraph("页码按当前版本排版生成；正文调整后可手动同步目录页码。")
    note.paragraph_format.space_before = Pt(8)
    for r in note.runs:
        r.font.name = "宋体"
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(110, 110, 110)
    anchor = insert = anchor
    anchor.addnext(note._p)

    doc.save(str(OUT))
    print(OUT)


if __name__ == "__main__":
    main()
