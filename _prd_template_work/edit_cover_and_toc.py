from copy import deepcopy
from io import BytesIO
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


SRC = Path(r"D:\AI\人工智能项目\智能排考小程序\智能排考小程序产品需求文档_融合优化版.docx")
OUT = Path(r"D:\AI\人工智能项目\智能排考小程序\智能排考小程序产品需求文档_融合优化版_封面目录优化.docx")


def add_position(anchor, tag, relative_from, offset):
    pos = OxmlElement(tag)
    pos.set("relativeFrom", relative_from)
    val = OxmlElement("wp:posOffset")
    val.text = str(offset)
    pos.append(val)
    anchor.append(pos)


def inline_to_page_anchor(inline):
    anchor = OxmlElement("wp:anchor")
    anchor.set("distT", "0")
    anchor.set("distB", "0")
    anchor.set("distL", "0")
    anchor.set("distR", "0")
    anchor.set("simplePos", "0")
    anchor.set("relativeHeight", "0")
    anchor.set("behindDoc", "1")
    anchor.set("locked", "0")
    anchor.set("layoutInCell", "1")
    anchor.set("allowOverlap", "1")

    simple = OxmlElement("wp:simplePos")
    simple.set("x", "0")
    simple.set("y", "0")
    anchor.append(simple)
    add_position(anchor, "wp:positionH", "page", 0)
    add_position(anchor, "wp:positionV", "page", 0)
    anchor.append(deepcopy(inline.extent))
    anchor.append(OxmlElement("wp:wrapNone"))
    anchor.append(deepcopy(inline.docPr))
    anchor.append(deepcopy(inline.graphic))
    return anchor


def add_toc_field(paragraph, entries):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    run._r.extend([begin, instr, separate])

    for index, entry in enumerate(entries):
        item_run = paragraph.add_run(entry)
        item_run.font.name = "宋体"
        item_run.font.size = Pt(11)
        if index < len(entries) - 1:
            item_run.add_break()

    end_run = paragraph.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run._r.append(end)


def insert_after(anchor, paragraph):
    anchor.addnext(paragraph._p)
    return paragraph._p


def main():
    OUT.write_bytes(SRC.read_bytes())
    doc = Document(str(OUT))

    # Reuse the existing cover image, but place it in the first-page header so it
    # sits behind the cover typography and does not affect document flow.
    image_para = next(p for p in doc.paragraphs if p._p.xpath('.//*[local-name()="blip"]'))
    image_rel_id = image_para._p.xpath('.//*[local-name()="blip"]')[0].get(qn("r:embed"))
    image_part = doc.part.related_parts[image_rel_id]
    section = doc.sections[0]
    section.different_first_page_header_footer = True
    section.header_distance = Inches(0)
    first_header = section.first_page_header
    header_para = first_header.paragraphs[0]
    header_para.paragraph_format.space_before = Pt(0)
    header_para.paragraph_format.space_after = Pt(0)
    header_para.add_run().add_picture(BytesIO(image_part.blob), width=Inches(8.27), height=Inches(11.69))
    header_inline = header_para._p.xpath('.//*[local-name()="inline"]')[0]
    header_inline.getparent().replace(header_inline, inline_to_page_anchor(header_inline))
    image_para._p.getparent().remove(image_para._p)

    # The existing cover ends with a page-break paragraph. Insert a readable,
    # updateable TOC immediately after it, then start the body on page 3.
    cover_break = next(p for p in doc.paragraphs if p._p.xpath('.//*[local-name()="br" and @w:type="page"]'))
    heading_titles = []
    for p in doc.paragraphs:
        if p.style.name in {"一级标题", "二级标题", "三级标题"} and p.text.strip():
            heading_titles.append(p.text.strip())

    title = doc.add_paragraph("目录", style="一级标题")
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(10)
    toc_para = doc.add_paragraph()
    toc_para.paragraph_format.space_after = Pt(0)
    add_toc_field(toc_para, heading_titles)
    note = doc.add_paragraph("提示：在 Word 中右键目录并选择“更新域”，可刷新页码和层级。")
    note.paragraph_format.space_before = Pt(10)
    note.paragraph_format.space_after = Pt(0)
    for run in note.runs:
        run.font.name = "宋体"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(102, 102, 102)
    next_page = doc.add_paragraph()
    next_page.add_run().add_break(WD_BREAK.PAGE)

    anchor = cover_break._p
    for para in (title, toc_para, note, next_page):
        anchor = insert_after(anchor, para)

    settings = doc.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")
    doc.save(str(OUT))
    print(OUT)


if __name__ == "__main__":
    main()
